import glob
import win32com.client
import os
import shutil
import extract_msg
from docx import Document
import time
import Tesseract_OCR_Extract
import pptx
from pptx import Presentation
from docx import Document
from docx.shared import Inches
import os
import aspose.pdf as pdf
import pdf2docx
import PyPDF2

def is_pdf_password_protected(pdf_path):
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            return reader.is_encrypted
    except Exception as e:
        print(f"An error occurred while checking the PDF: {e}")
        return False

def PDF_to_Word(BasicPath, filename):
    pdfs_path    = BasicPath + "\\"  + filename
    word         = win32com.client.Dispatch("Word.Application")
    
    start_time = time.time()
    if is_pdf_password_protected(pdfs_path):
        try:
            try:
                print("Convertion Failure     : Password Protected")
                print("Trying Again...")

                in_file = os.path.abspath(pdfs_path)
                wb = word.Documents.Open(in_file)
                out_file = os.path.abspath(pdfs_path[0:-4] + ".docx")
                wb.SaveAs2(out_file, FileFormat=16)
                wb.Close()
                end_time = time.time()
                elapsed_time = round(((end_time - start_time) / 60),2)
                print("Converted PDFs to Docx : ")
            except:
                print("Convertion Failure     : ")
                print("It is Scanned Copy - Trying Tesseract OCR ...")
                Tesseract_OCR_Extract.Convert_Scanned_PDF_to_Docx(pdfs_path, pdfs_path[0:-5] + ".docx")
                end_time = time.time()
                elapsed_time = round(((end_time - start_time) / 60),2)
                print("Converted PDFs to Docx : ")
        except:
            print("Convertion Failure     : During Final Try")
            print("Ignoring...")
    else:
        try:
            try:
                
                in_file = os.path.abspath(pdfs_path)
                
                wb = word.Documents.Open(in_file)
                out_file = os.path.abspath(pdfs_path[0:-4] + ".docx")
                wb.SaveAs2(out_file, FileFormat=16)
                wb.Close()
                end_time = time.time()
                elapsed_time = round(((end_time - start_time) / 60),2)
                print("Converted PDFs to Docx : ")
            except:
                print("Convertion Failure     : ")
                print("It is Scanned Copy - Trying Tesseract OCR ...")
                Tesseract_OCR_Extract.Convert_Scanned_PDF_to_Docx(pdfs_path, pdfs_path[0:-5] + ".docx")
                end_time = time.time()
                elapsed_time = round(((end_time - start_time) / 60),2)
                print("Converted PDFs to Docx : ")
        except:
            print("Convertion Failure     : During Final Try")
            print("Ignoring...")
    
    print(" ")
                
    word.Quit()
    
def Docx_to_Text(BasicPath, filename):
    
    word_path    = BasicPath + "\\" + filename
    doc          = Document(word_path)
    full_text = []

    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = ' '.join(paragraph.text for paragraph in cell.paragraphs)
                row_text.append(cell_text)
            full_text.append('\t'.join(row_text))  # Join each cell text with a tab

    text = '\n'.join(full_text)
    return text
    



